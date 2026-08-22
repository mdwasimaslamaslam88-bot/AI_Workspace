from __future__ import annotations

import csv
import io
import logging
import re
import unicodedata
import zipfile
from dataclasses import dataclass
from pathlib import PurePosixPath

from docx import Document as DocxDocument
from pypdf import PdfReader


MAX_DOCUMENT_BYTES = 1_048_576
MAX_DOCUMENT_CHARACTERS = 1_000_000
MAX_DOCUMENT_UNITS = 10_000
MAX_DOCUMENT_CHUNKS = 1_024
MAX_PDF_PAGES = 200
MAX_CSV_COLUMNS = 256
MAX_CSV_FIELD_CHARACTERS = 32_768
MAX_DOCX_ARCHIVE_MEMBERS = 512
MAX_DOCX_UNCOMPRESSED_BYTES = 8_388_608
MAX_DOCX_COMPRESSION_RATIO = 200
DOCUMENT_CHUNK_CHARACTERS = 1_200
DOCUMENT_CHUNK_OVERLAP_CHARACTERS = 200
DOCX_MEDIA_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
SUPPORTED_DOCUMENT_MEDIA_TYPES = frozenset(
    {"text/plain", "text/csv", "application/pdf", DOCX_MEDIA_TYPE}
)

logging.getLogger("pypdf").setLevel(logging.ERROR)


class DocumentParseError(ValueError):
    """Uploaded bytes are not a supported, safe, parseable document."""


class DocumentTooLargeError(DocumentParseError):
    """A document exceeds a fixed parser or extracted-text bound."""


@dataclass(frozen=True, slots=True)
class ParsedUnit:
    content: str
    provenance_kind: str
    page_number: int | None = None
    row_start: int | None = None
    row_end: int | None = None
    section: str | None = None


@dataclass(frozen=True, slots=True)
class DocumentChunkDraft:
    ordinal: int
    content: str
    provenance_kind: str
    page_number: int | None
    row_start: int | None
    row_end: int | None
    section: str | None


def _normalized_content(value: str) -> str:
    if "\x00" in value:
        raise DocumentParseError("document text contains unsupported control data")
    normalized = unicodedata.normalize("NFC", value)
    return re.sub(r"\s+", " ", normalized).strip()


def _validate_units(units: list[ParsedUnit]) -> tuple[ParsedUnit, ...]:
    if len(units) > MAX_DOCUMENT_UNITS:
        raise DocumentTooLargeError("document contains too many text units")
    character_count = 0
    normalized: list[ParsedUnit] = []
    for unit in units:
        content = _normalized_content(unit.content)
        if not content:
            continue
        character_count += len(content)
        if character_count > MAX_DOCUMENT_CHARACTERS:
            raise DocumentTooLargeError("document text is too large")
        normalized.append(
            ParsedUnit(
                content=content,
                provenance_kind=unit.provenance_kind,
                page_number=unit.page_number,
                row_start=unit.row_start,
                row_end=unit.row_end,
                section=(
                    _normalized_content(unit.section)[:255]
                    if unit.section
                    else None
                ),
            )
        )
    if not normalized:
        raise DocumentParseError("document contains no readable text")
    return tuple(normalized)


def _parse_text(data: bytes) -> tuple[ParsedUnit, ...]:
    try:
        content = data.decode("utf-8", errors="strict")
    except UnicodeDecodeError as exc:
        raise DocumentParseError("text document must be UTF-8") from exc
    return _validate_units([ParsedUnit(content, "text")])


def _parse_csv(data: bytes) -> tuple[ParsedUnit, ...]:
    try:
        content = data.decode("utf-8-sig", errors="strict")
    except UnicodeDecodeError as exc:
        raise DocumentParseError("CSV document must be UTF-8") from exc
    units: list[ParsedUnit] = []
    try:
        reader = csv.reader(io.StringIO(content, newline=""), strict=True)
        for row_number, row in enumerate(reader, start=1):
            if row_number > MAX_DOCUMENT_UNITS:
                raise DocumentTooLargeError("CSV contains too many rows")
            if len(row) > MAX_CSV_COLUMNS:
                raise DocumentTooLargeError("CSV contains too many columns")
            if any(len(field) > MAX_CSV_FIELD_CHARACTERS for field in row):
                raise DocumentTooLargeError("CSV field is too large")
            units.append(
                ParsedUnit(
                    " | ".join(row),
                    "row",
                    row_start=row_number,
                    row_end=row_number,
                )
            )
    except (csv.Error, UnicodeError) as exc:
        raise DocumentParseError("CSV document is malformed") from exc
    return _validate_units(units)


def _parse_pdf(data: bytes) -> tuple[ParsedUnit, ...]:
    try:
        reader = PdfReader(io.BytesIO(data), strict=True)
        if reader.is_encrypted:
            raise DocumentParseError("encrypted PDFs are not supported")
        if len(reader.pages) > MAX_PDF_PAGES:
            raise DocumentTooLargeError("PDF contains too many pages")
        units = [
            ParsedUnit(page.extract_text() or "", "page", page_number=index)
            for index, page in enumerate(reader.pages, start=1)
        ]
    except DocumentParseError:
        raise
    except Exception as exc:
        raise DocumentParseError("PDF document is malformed") from exc
    return _validate_units(units)


def _validate_docx_package(data: bytes) -> None:
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            members = archive.infolist()
            if len(members) > MAX_DOCX_ARCHIVE_MEMBERS:
                raise DocumentTooLargeError("DOCX package contains too many members")
            names: set[str] = set()
            total_uncompressed = 0
            for member in members:
                path = PurePosixPath(member.filename.replace("\\", "/"))
                if path.is_absolute() or ".." in path.parts:
                    raise DocumentParseError("DOCX package path is unsafe")
                if member.flag_bits & 0x1:
                    raise DocumentParseError("encrypted DOCX packages are not supported")
                total_uncompressed += member.file_size
                if total_uncompressed > MAX_DOCX_UNCOMPRESSED_BYTES:
                    raise DocumentTooLargeError("DOCX package expands beyond its bound")
                if (
                    (member.file_size > 0 and member.compress_size == 0)
                    or (
                        member.compress_size > 0
                        and member.file_size
                        > member.compress_size * MAX_DOCX_COMPRESSION_RATIO
                    )
                ):
                    raise DocumentTooLargeError("DOCX compression ratio is unsafe")
                names.add(member.filename)
            if "[Content_Types].xml" not in names or "word/document.xml" not in names:
                raise DocumentParseError("DOCX package is incomplete")
    except DocumentParseError:
        raise
    except (OSError, zipfile.BadZipFile, zipfile.LargeZipFile) as exc:
        raise DocumentParseError("DOCX document is malformed") from exc


def _parse_docx(data: bytes) -> tuple[ParsedUnit, ...]:
    _validate_docx_package(data)
    try:
        document = DocxDocument(io.BytesIO(data))
        units: list[ParsedUnit] = []
        current_section: str | None = None
        for paragraph in document.paragraphs:
            if len(units) >= MAX_DOCUMENT_UNITS:
                raise DocumentTooLargeError("DOCX contains too many paragraphs")
            text = paragraph.text
            style_name = paragraph.style.name if paragraph.style is not None else ""
            if style_name.startswith("Heading") and text.strip():
                current_section = text
            units.append(
                ParsedUnit(
                    text,
                    "section" if current_section else "text",
                    section=current_section,
                )
            )
        for table in document.tables:
            for row_number, row in enumerate(table.rows, start=1):
                if len(units) >= MAX_DOCUMENT_UNITS:
                    raise DocumentTooLargeError("DOCX contains too many text units")
                units.append(
                    ParsedUnit(
                        " | ".join(cell.text for cell in row.cells),
                        "row",
                        row_start=row_number,
                        row_end=row_number,
                        section=current_section,
                    )
                )
    except DocumentParseError:
        raise
    except Exception as exc:
        raise DocumentParseError("DOCX document is malformed") from exc
    return _validate_units(units)


def parse_document(
    data: bytes,
    media_type: str,
) -> tuple[ParsedUnit, ...]:
    if not isinstance(data, bytes):
        raise TypeError("document data must be bytes")
    if not 1 <= len(data) <= MAX_DOCUMENT_BYTES:
        raise DocumentTooLargeError("document byte size is outside the parser bound")
    if media_type not in SUPPORTED_DOCUMENT_MEDIA_TYPES:
        raise DocumentParseError("document media type is unsupported")
    parser = {
        "text/plain": _parse_text,
        "text/csv": _parse_csv,
        "application/pdf": _parse_pdf,
        DOCX_MEDIA_TYPE: _parse_docx,
    }[media_type]
    return parser(data)


def chunk_document(units: tuple[ParsedUnit, ...]) -> tuple[DocumentChunkDraft, ...]:
    drafts: list[DocumentChunkDraft] = []
    for unit in units:
        start = 0
        while start < len(unit.content):
            maximum_end = min(start + DOCUMENT_CHUNK_CHARACTERS, len(unit.content))
            end = maximum_end
            if maximum_end < len(unit.content):
                split = unit.content.rfind(
                    " ",
                    start + DOCUMENT_CHUNK_CHARACTERS // 2,
                    maximum_end,
                )
                if split > start:
                    end = split
            chunk = unit.content[start:end].strip()
            if chunk:
                drafts.append(
                    DocumentChunkDraft(
                        ordinal=len(drafts) + 1,
                        content=chunk,
                        provenance_kind=unit.provenance_kind,
                        page_number=unit.page_number,
                        row_start=unit.row_start,
                        row_end=unit.row_end,
                        section=unit.section,
                    )
                )
                if len(drafts) > MAX_DOCUMENT_CHUNKS:
                    raise DocumentTooLargeError("document creates too many chunks")
            if end >= len(unit.content):
                break
            next_start = max(
                end - DOCUMENT_CHUNK_OVERLAP_CHARACTERS,
                start + 1,
            )
            while next_start < end and unit.content[next_start].isspace():
                next_start += 1
            start = next_start
    if not drafts:
        raise DocumentParseError("document contains no indexable text")
    return tuple(drafts)
